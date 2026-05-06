from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Optional

from baudoku_api.reports.image_processing import process_report_image
from baudoku_api.reports.letterhead import (
    BBA_BLUE,
    BBA_TEXT,
    CONTACT_COLUMNS,
    logo_bytes_from_template,
)
from baudoku_api.reports.report_content import (
    ReportContent,
    ReportImage,
    build_report_content,
)


ImageLoader = Callable[[str], bytes]

DEFAULT_TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "templates" / "bba_briefbogen.docx"
PHOTO_MAX_LONG_EDGE_PX = 1800
PHOTO_MAX_WIDTH_INCHES = 4.5
PHOTO_MAX_HEIGHT_INCHES = 3.8
PLAN_MAX_LONG_EDGE_PX = 2400
PLAN_MAX_WIDTH_INCHES = 6.5
PLAN_MAX_HEIGHT_INCHES = 7.2


class ReportDocxBuilderError(Exception):
    """Raised when a DOCX report cannot be built."""


class ReportDocxBuilder:
    def __init__(
        self,
        template_path: Optional[object] = None,
        image_loader: Optional[ImageLoader] = None,
    ) -> None:
        self.template_path = Path(str(template_path)).expanduser() if template_path else None
        self.image_loader = image_loader

    def build(
        self,
        project: dict[str, Any],
        defects: list[dict[str, Any]],
        general_findings: list[dict[str, Any]],
        project_conclusion: Optional[dict[str, Any]],
        plans: Optional[list[dict[str, Any]]] = None,
    ) -> bytes:
        document = self._document_from_template()
        self._prepare_document(document)
        content = build_report_content(project, defects, general_findings, project_conclusion, plans)
        self._build_content(document, content)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _document_from_template(self) -> Any:
        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover - optional dependency surface
            raise ReportDocxBuilderError("python-docx ist nicht installiert.") from exc

        template_path = self._resolved_template_path()
        try:
            return Document(str(template_path))
        except Exception as exc:
            raise ReportDocxBuilderError(
                f"Briefbogen-Template konnte nicht geladen werden: {template_path}"
            ) from exc

    def _resolved_template_path(self) -> Path:
        candidates = []
        if self.template_path is not None:
            candidates.append(self.template_path)
        candidates.append(DEFAULT_TEMPLATE_PATH)

        for candidate in candidates:
            if candidate.exists():
                return candidate

        raise ReportDocxBuilderError(
            f"Briefbogen-Template fehlt. Erwarteter Fallback: {DEFAULT_TEMPLATE_PATH}"
        )

    def _prepare_document(self, document: Any) -> None:
        self._clear_body(document)
        self._configure_page(document)
        self._configure_styles(document)
        self._configure_letterhead(document)

    def _build_content(self, document: Any, content: ReportContent) -> None:
        document.add_paragraph(content.title, style="Title")
        self._add_project_header(document, content)
        self._add_general_findings(document, content)
        self._add_defects(document, content)
        self._add_plans(document, content)
        if content.conclusion:
            document.add_heading("Fazit", level=1)
            document.add_paragraph(content.conclusion)

    def _clear_body(self, document: Any) -> None:
        from docx.oxml.ns import qn

        body = document._body._element
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    def _configure_page(self, document: Any) -> None:
        from docx.shared import Mm

        section = document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(17.5)
        section.right_margin = Mm(7.5)
        section.top_margin = Mm(42)
        section.bottom_margin = Mm(36)
        section.header_distance = Mm(8)
        section.footer_distance = Mm(7.2)
        section.different_first_page_header_footer = True

    def _configure_styles(self, document: Any) -> None:
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Pt, RGBColor

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08

        title = document.styles["Title"]
        title.font.name = "Arial"
        title.font.size = Pt(18)
        title.font.bold = True
        title.font.color.rgb = RGBColor.from_string(BBA_TEXT)
        title.paragraph_format.space_after = Pt(12)

        for style_name, size, before, after in (
            ("Heading 1", 14, 12, 6),
            ("Heading 2", 12, 9, 4),
            ("Heading 3", 10.5, 7, 3),
        ):
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.bold = True
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(BBA_TEXT)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)

        try:
            caption = document.styles["BBA Caption"]
        except KeyError:
            caption = document.styles.add_style("BBA Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "Arial"
        caption.font.size = Pt(8)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor(85, 85, 85)
        caption.paragraph_format.space_after = Pt(8)

    def _configure_letterhead(self, document: Any) -> None:
        for header in (document.sections[0].first_page_header, document.sections[0].header):
            self._build_header(header)
        for footer in (document.sections[0].first_page_footer, document.sections[0].footer):
            self._build_footer(footer)

    def _build_header(self, header: Any) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt

        _clear_block_container(header)
        logo = logo_bytes_from_template(self._resolved_template_path())
        logo_paragraph = header.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_paragraph.paragraph_format.space_after = Pt(5)
        run = logo_paragraph.add_run()
        if logo:
            run.add_picture(BytesIO(logo), width=Mm(52.6))
        else:
            run.bold = True
            run.text = "BBA GmbH"

        line = header.add_paragraph()
        line.paragraph_format.space_after = Pt(0)
        _set_paragraph_bottom_border(line, BBA_BLUE, size="4")

    def _build_footer(self, footer: Any) -> None:
        from docx.shared import Mm, Pt

        _clear_block_container(footer)
        line = footer.add_paragraph()
        line.paragraph_format.space_after = Pt(3)
        _set_paragraph_bottom_border(line, BBA_BLUE, size="4")

        table = footer.add_table(rows=1, cols=4, width=Mm(185))
        table.autofit = False
        _clear_table_borders(table)
        widths = [Mm(46.25)] * 4
        for column_index, (heading, lines) in enumerate(CONTACT_COLUMNS):
            cell = table.cell(0, column_index)
            _set_cell_width(cell, widths[column_index])
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(heading)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(7)
            for line_text in lines:
                paragraph.add_run().add_break()
                line_run = paragraph.add_run(line_text)
                line_run.font.name = "Arial"
                line_run.font.size = Pt(7)

    def _add_project_header(self, document: Any, content: ReportContent) -> None:
        from docx.shared import Inches, Pt

        table = document.add_table(rows=len(content.project_fields), cols=2)
        table.style = "Table Grid"
        table.autofit = False
        _set_table_width(table, Inches(7.28))
        _set_table_borders(table, color="D6DCE2", size="4")
        for row, (label, value) in zip(table.rows, content.project_fields):
            _set_cell_width(row.cells[0], Inches(1.55))
            _set_cell_width(row.cells[1], Inches(5.73))
            row.cells[0].text = label
            row.cells[1].text = value or "Nicht angegeben"
            for paragraph in row.cells[0].paragraphs:
                paragraph.runs[0].bold = True
            for cell_index, cell in enumerate(row.cells):
                _set_cell_shading(cell, "F3F7FA" if cell_index == 0 else "FFFFFF")
                _set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)

        document.add_paragraph()

    def _add_general_findings(self, document: Any, content: ReportContent) -> None:
        document.add_heading("Allgemeine Feststellungen", level=1)
        if not content.general_findings:
            document.add_paragraph("Es wurden noch keine allgemeinen Feststellungen erfasst.")
            return

        for finding in content.general_findings:
            paragraph = document.add_paragraph(finding, style="List Bullet")
            paragraph.paragraph_format.space_after = 0

    def _add_defects(self, document: Any, content: ReportContent) -> None:
        document.add_heading("Mängel und Hinweise", level=1)
        if not content.defects:
            document.add_paragraph("Es wurden noch keine Mängel oder Hinweise erfasst.")
            return

        for defect in content.defects:
            document.add_heading(f"{defect.label} - {defect.kind}", level=2)
            if defect.trade:
                document.add_paragraph(f"Gewerk: {defect.trade}")
            if defect.category:
                document.add_paragraph(f"Kategorie: {defect.category}")
            document.add_paragraph(defect.description or "Ohne Beschreibung")
            if defect.photos:
                document.add_heading("Fotos", level=3)
            for photo in defect.photos:
                self._add_image_block(document, photo)

    def _add_plans(self, document: Any, content: ReportContent) -> None:
        if not content.plans:
            return

        document.add_heading("Planverortung", level=1)
        for plan in content.plans:
            document.add_heading(plan.name, level=2)
            if plan.storage_path:
                self._add_image_block(
                    document,
                    ReportImage(
                        storage_path=plan.storage_path,
                        caption=f"Plan {plan.name}: markierte Verortung",
                        label="",
                        kind="plan",
                    ),
                )
            if plan.render_error:
                document.add_paragraph(plan.render_error)

    def _add_image_block(self, document: Any, image: ReportImage) -> None:
        if not self.image_loader or not image.storage_path:
            return

        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches

        max_long_edge = PLAN_MAX_LONG_EDGE_PX if image.kind == "plan" else PHOTO_MAX_LONG_EDGE_PX
        max_width = PLAN_MAX_WIDTH_INCHES if image.kind == "plan" else PHOTO_MAX_WIDTH_INCHES
        max_height = PLAN_MAX_HEIGHT_INCHES if image.kind == "plan" else PHOTO_MAX_HEIGHT_INCHES
        try:
            processed = process_report_image(
                self.image_loader(image.storage_path),
                max_long_edge_px=max_long_edge,
                max_width_inches=max_width,
                max_height_inches=max_height,
            )
        except Exception:
            document.add_paragraph(f"Bild konnte nicht eingebettet werden: {image.storage_path}")
            return

        picture_paragraph = document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.keep_with_next = True
        run = picture_paragraph.add_run()
        run.add_picture(
            BytesIO(processed.content),
            width=Inches(processed.display_width_inches),
            height=Inches(processed.display_height_inches),
        )
        caption_paragraph = document.add_paragraph(style="BBA Caption")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.keep_together = True
        caption_paragraph.add_run(f"{image.label + ': ' if image.label else ''}{image.caption}")


def _clear_block_container(container: Any) -> None:
    for paragraph in list(container.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(container.tables):
        table._element.getparent().remove(table._element)


def _set_paragraph_bottom_border(paragraph: Any, color: str, size: str = "4") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def _clear_table_borders(table: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _set_table_borders(table: Any, color: str, size: str = "4") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    tbl_pr.append(borders)


def _set_table_width(table: Any, width: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width.twips))
    tbl_w.set(qn("w:type"), "dxa")


def _set_cell_width(cell: Any, width: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width.twips))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = width


def _set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, **margins: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in margins.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
